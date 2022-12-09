import time
import os
import sys
from itertools import zip_longest
import xlwings as xw
from docx import Document
from docx.shared import Inches, Pt

# -- Global declarations
start_time = time.time()
os.chdir(sys.path[0])
doc = Document()
xw.Book("ScopeData.xlsx").set_mock_caller()
wb = xw.Book.caller()

def string_conversion(old_list):
    converted_list = []
    for value in old_list:
        if type(value) == float:
            converted_list.append(str(int(value)))
        else:
            converted_list.append(str(value))
    return converted_list

def info_dictionary():
    raw_keys = string_conversion(wb.sheets[1].range('A1').expand('down').value)
    raw_values = wb.sheets[1].range('B1').expand('down').value
    return dict(zip(raw_keys, raw_values))
    
def scrape_excel():

    def flexed_range(col):
        flag = wb.sheets[0].range('F' + str(wb.sheets[0].cells.last_cell.row)).end('up').row
        range = str(col) + '2:' + str(col) + str(flag)
        return range

    a_values = string_conversion(wb.sheets[0].range(flexed_range('A')).value)
    b_values = string_conversion(wb.sheets[0].range(flexed_range('B')).value)
    c_values = string_conversion(wb.sheets[0].range(flexed_range('C')).value)
    d_values = string_conversion(wb.sheets[0].range(flexed_range('D')).value)
    e_values = string_conversion(wb.sheets[0].range(flexed_range('E')).value)

    data = list(zip_longest(a_values, 
                            b_values, 
                            c_values, 
                            d_values, 
                            e_values, 
                            fillvalue='None'))
    return data

# -- Needs work to integrate
def low_voltage():
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    format = paragraph.paragraph_format
    font = run.font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    format.space_before = Pt(0)
    format.space_after = Pt(0)
    run.add_text('Low Voltage Electrical Option:')
    font.bold = True
    font.underline = True
    
    paragraph = doc.add_paragraph(
        'Supply a conduit layout drawing for devices called out in automation scope.', 
        style='List Bullet'
        )
    paragraph = doc.add_paragraph(
        'Provide and install wet rated, low voltage wire in existing conduit raceways (by others).', 
        style='List Bullet'
        )
    paragraph = doc.add_paragraph(
        'Final wiring and termination of low voltage devices.', 
        style='List Bullet'
        )
    paragraph = doc.add_paragraph(
        'Supply low voltage electrical permit & schedule inspection with '
        'local municipalities when required.', 
        style='List Bullet'
        )
    paragraph = doc.add_paragraph(
        'Test system.', style='List Bullet'
        )

def format_scopes(header, work, qty, info, bullet):
    raw_dict = info_dictionary()

    # -- Declaring format
    paragraph = doc.add_paragraph()
    format = paragraph.paragraph_format
    tab_stop = format.tab_stops
    tab_stop.add_tab_stop(Inches(0.5))
    tab_stop.add_tab_stop(Inches(0.8))
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    format.space_before = Pt(0)
    format.space_after = Pt(0)


    def delete_paragraph(paragraph):
        line = paragraph._element
        line.getparent().remove(line)
        line._p = line._element = None
    

    
    if header == 'None':
        pass
    else:
        run = paragraph.add_run(header + ':')
        paragraph = doc.add_paragraph()
        format = paragraph.paragraph_format
        font = run.font
        font.bold = True
        font.underline = True
        tab_stop = format.tab_stops
        tab_stop.add_tab_stop(Inches(0.5))
        tab_stop.add_tab_stop(Inches(0.8))


    if work == 'None':
        run = paragraph.add_run()
        run.add_tab()
    else:             
        run = paragraph.add_run(work + ':')
        run.add_tab()


    if header == 'None' and work == 'None' and qty == 'None' and info == 'None':
        delete_paragraph(paragraph)
    elif qty == 'None':
        run = paragraph.add_run()
        run.add_text(' -')
        run.add_tab()
    else:
        run = paragraph.add_run()
        run.add_text(qty + ')')
        run.add_tab()


    if info == 'None':
        pass
    elif info in raw_dict.keys():
        run = paragraph.add_run()
        run.add_text(raw_dict[info])
    else:
        run = paragraph.add_run()
        run.add_text(info)


    # -- Bullet point
    if bullet == 'None':
        pass
    else: 
        paragraph = doc.add_paragraph(bullet, style='List Bullet')
        format = paragraph.paragraph_format
        format.space_before = Pt(0)
        format.space_after = Pt(0)

def create_word():
    data = scrape_excel()
    for entry in data:
        tuple = entry
        header, work, qty, info, bullet = tuple
        format_scopes(header, work, qty, info, bullet)

def format_wen():
    
    # -- Not working with only 1 entry in H, I, J columns.
    w_data = string_conversion(wb.sheets[0].range('H2').expand('down').value)
    e_data = string_conversion(wb.sheets[0].range('I2').expand('down').value)
    n_data = string_conversion(wb.sheets[0].range('J2').expand('down').value)

    w_numbers = ['W' + item for item in w_data]
    e_numbers = ['EX' + item for item in e_data]
    n_numbers = ['N' + item for item in n_data]
  
    # -- Dictionary
    raw_keys = wb.sheets[2].range('A1').expand('down').value
    raw_values = wb.sheets[2].range('B1').expand('down').value
    raw_dict = dict(zip(raw_keys, raw_values))
  

    # -- Dictionaries based on autocorrect code/values
    warranties = [raw_dict[key] for key in w_numbers]
    exclusions = [raw_dict[key] for key in e_numbers]  
    notes = [raw_dict[key] for key in n_numbers]
    

    # -- Add and format to Word
    p = doc.add_paragraph()
    run = p.add_run('Warranties:')
    run.bold = True
    run.underline = True
    for entry in warranties:
        p = doc.add_paragraph(entry)
        p.style = 'List Bullet'

    p = doc.add_paragraph()
    run = p.add_run('Exclusions:')
    run.bold = True
    run.underline = True
    for entry in exclusions:
        p = doc.add_paragraph(entry)
        p.style = 'List Bullet'

    p = doc.add_paragraph()
    run = p.add_run('Notes:')
    run.bold = True
    run.underline = True
    for entry in notes:
        p = doc.add_paragraph(entry)
        p.style = 'List Bullet'

def start_program():
    create_word()
    format_wen()
    doc.save('Rendered.docx')
    os.system('start Rendered.docx')

start_program()


print('Process finished in %s seconds ' % (time.time() - start_time))